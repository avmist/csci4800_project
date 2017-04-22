import requests
import json
import getpass
import docx
import pprint

# initialize login JSON object for requesting token
def login(ip,un,pw,port):
    base_url = 'https://' + ip + ':' + port
    login = {'username':un,'password':pw}

# get token with login
def get_token(login,base_url):
    resp = requests.post(base_url+'/session',data=login,verify=False)
    return resp.json()['token']

# get a list of all scans
def get_scans(login,base_url):
    scan_names = []
    token = get_token(login,base_url)
    headers = {'X-Cookie': 'token=' + token, 'content-type': 'application/json'}
    scans = json.loads(requests.get(base_url + '/scans', data=None, headers=headers, verify=False).content)['scans']
    for scan in scans:
        scan_names.append(scan['name'])
    return scan_names

# given a scan name, get a list of all items in that scan
def get_scan_data(login,base_url,name):
    scan_items = []
    token = get_token(login,base_url)
    headers = {'X-Cookie': 'token=' + token, 'content-type': 'application/json'}
    scans = requests.get(base_url + '/scans', data=None, headers=headers, verify=False).json()['scans']
    for scan in scans:
        if scan['name'] == name:
            id = scan['id']
    scan_data = json.loads(requests.get(base_url + '/scans/' + str(id), data=None, headers=headers, verify=False).content)
    return scan_data.keys()

def createDocument(login,base_url,chosen_scan,opts,loc):
    if not loc.endswith('.docx'):
        loc += '.docx'
    print loc
    token = get_token(login,base_url)
    headers = {'X-Cookie': 'token=' + token, 'content-type': 'application/json'}
    scans = requests.get(base_url + '/scans', data=None, headers=headers, verify=False).json()['scans']
    for scan in scans:
        if scan['name'] == chosen_scan:
            id = scan['id']
    scan_data = json.loads(requests.get(base_url + '/scans/' + str(id), data=None, headers=headers, verify=False).content)
    doc = docx.Document()
    for opt in opts:
        doc.add_heading(opt,1)
        jsonObj = json.dumps(scan_data[opt],indent=4).replace(',','').replace('[','').replace(']','').replace('{','').replace('}','').replace('"','')
        print jsonObj
        for i in jsonObj.splitlines():
            doc.add_paragraph(i)
    doc.save(loc)
